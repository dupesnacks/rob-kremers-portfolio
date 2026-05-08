#!/usr/bin/env python3
"""Convert TVRN tweets JSONL to Word document (.docx)"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def jsonl_to_docx(jsonl_file, output_file):
    doc = Document()
    
    # Title
    title = doc.add_heading('@tvRN20 Tweet Archive', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {Path(jsonl_file).stat().st_mtime}\n").font.size = Pt(10)
    meta.add_run(f"Source: {jsonl_file}").font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # Load and count
    tweets = []
    with open(jsonl_file) as f:
        for line in f:
            tweets.append(json.loads(line))
    
    doc.add_heading(f'{len(tweets)} Tweets', level=1)
    doc.add_paragraph()
    
    # Add each tweet
    for i, tweet in enumerate(tweets, 1):
        # Tweet number and date
        p = doc.add_paragraph()
        p.add_run(f"#{i} — {tweet['created_at']}").font.size = Pt(9)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        
        # Tweet text
        p = doc.add_paragraph(tweet['text'])
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        
        # Engagement metrics (smaller)
        metrics = tweet.get('public_metrics', {})
        if metrics:
            p = doc.add_paragraph()
            p.add_run(f"❤️ {metrics.get('like_count', 0)} | 🔄 {metrics.get('retweet_count', 0)} | 💬 {metrics.get('reply_count', 0)}").font.size = Pt(8)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(8)
    
    # Save
    doc.save(output_file)
    print(f"✅ Saved {len(tweets)} tweets to {output_file}")
    print(f"File size: {Path(output_file).stat().st_size / 1024 / 1024:.1f} MB")

if __name__ == "__main__":
    jsonl = Path.home() / "clawd" / "memory" / "tvrn-tweets.jsonl"
    docx = Path.home() / "clawd" / "TVRN_Tweets_Archive.docx"
    
    jsonl_to_docx(jsonl, docx)
