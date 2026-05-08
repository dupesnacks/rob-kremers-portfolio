#!/usr/bin/env python3
"""
Fetch Raul's (@unfairmarket) tweets from the last year via X API v2.
Output: raul_tweets.json + raul_tweets.docx
"""

import os
import json
import requests
from datetime import datetime, timedelta
from pathlib import Path
from urllib.parse import unquote
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load token from .env.local
env_path = Path("/Users/rk/clawd/.env.local")
token_encoded = None
if env_path.exists():
    with open(env_path) as f:
        for line in f:
            if "X_BEARER_TOKEN=" in line:
                token_encoded = line.split("=", 1)[1].strip()
                break

if not token_encoded:
    print("ERROR: X_BEARER_TOKEN not found in .env.local")
    exit(1)

# URL decode the token
token = unquote(token_encoded)
print(f"✓ Token loaded (length: {len(token)})")

# X API v2 headers
headers = {
    "Authorization": f"Bearer {token}",
    "User-Agent": "TwitterBot/1.0"
}

# Step 1: Get user ID for @unfairmarket
print("\n[1/3] Fetching user ID for @unfairmarket...")
user_response = requests.get(
    "https://api.twitter.com/2/users/by/username/unfairmarket",
    headers=headers,
    params={"user.fields": "created_at,public_metrics"}
)

if user_response.status_code != 200:
    print(f"ERROR: Failed to get user info: {user_response.status_code}")
    print(user_response.text)
    exit(1)

user_data = user_response.json()
user_id = user_data["data"]["id"]
user_name = user_data["data"]["name"]
print(f"✓ Found user: {user_name} (ID: {user_id})")

# Step 2: Fetch tweets from last year with pagination
print("\n[2/3] Fetching tweets from last 365 days...")

tweets = []
start_time = (datetime.utcnow() - timedelta(days=365)).strftime("%Y-%m-%dT%H:%M:%SZ")
print(f"  Start time: {start_time}")

pagination_token = None
page = 1

while True:
    print(f"  • Page {page}...", end="", flush=True)
    
    params = {
        "max_results": 100,
        "start_time": start_time,
        "tweet.fields": "created_at,public_metrics,author_id,conversation_id",
        "expansions": "author_id",
        "user.fields": "username,name,created_at,public_metrics",
        "exclude": "retweets"  # Only original tweets, no RTs
    }
    
    if pagination_token:
        params["pagination_token"] = pagination_token
    
    response = requests.get(
        f"https://api.twitter.com/2/users/{user_id}/tweets",
        headers=headers,
        params=params
    )
    
    if response.status_code != 200:
        print(f"\nERROR at page {page}: {response.status_code}")
        print(response.text)
        break
    
    data = response.json()
    
    if "data" not in data or len(data["data"]) == 0:
        print(" (done)")
        break
    
    tweets.extend(data["data"])
    print(f" {len(data['data'])} tweets")
    
    # Check for pagination
    if "meta" in data and "next_token" in data["meta"]:
        pagination_token = data["meta"]["next_token"]
        page += 1
    else:
        break

print(f"\n✓ Total tweets fetched: {len(tweets)}")

# Step 3: Structure and save JSON
print("\n[3/3] Processing and saving output...")

structured_tweets = []
for tweet in tweets:
    structured_tweets.append({
        "id": tweet["id"],
        "created_at": tweet["created_at"],
        "text": tweet["text"],
        "public_metrics": tweet.get("public_metrics", {}),
        "conversation_id": tweet.get("conversation_id", "")
    })

# Sort by date descending (newest first)
structured_tweets.sort(key=lambda x: x["created_at"], reverse=True)

# Save JSON
json_path = Path("/Users/rk/clawd/raul_tweets.json")
with open(json_path, "w") as f:
    json.dump({
        "user": {
            "name": user_name,
            "handle": "@unfairmarket",
            "id": user_id,
            "fetched_at": datetime.utcnow().isoformat()
        },
        "tweets": structured_tweets,
        "count": len(structured_tweets),
        "period": f"Last 365 days (from {start_time})"
    }, f, indent=2)

print(f"✓ JSON saved: {json_path}")

# Generate DOCX
print("  Generating DOCX...")
doc = Document()

# Title
title = doc.add_heading(f"@{user_name} (@unfairmarket) — Tweet Archive", 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Metadata
meta = doc.add_paragraph()
meta.add_run(f"Fetched: ").bold = True
meta.add_run(f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n")
meta.add_run(f"Total Tweets: ").bold = True
meta.add_run(f"{len(structured_tweets)}\n")
meta.add_run(f"Period: ").bold = True
meta.add_run(f"Last 365 days\n")

# Table of Contents by theme
doc.add_heading("Quick Theme Index", 1)
themes = {
    "Outfits & Numbers": ["outfit", "angle", "MA", "period", "REPEATER", "SVIX", "NVDA"],
    "TSLA & Leverage": ["TSLA", "TSLQ", "TSLL", "TSLR", "TSLT", "TSLZ"],
    "BYND & Squeezes": ["BYND", "squeeze", "pressure", "breakout", "support"],
    "Market Calls": ["bullish", "bearish", "bias", "regime", "Fed"],
    "BTC & Crypto": ["Bitcoin", "BTC", "crypto", "Coinbase"],
    "Institutional": ["institutional", "conviction", "L/S", "hedge", "position"]
}

for theme, keywords in themes.items():
    count = sum(1 for t in structured_tweets if any(kw.lower() in t["text"].lower() for kw in keywords))
    doc.add_paragraph(f"{theme}: {count} tweets", style='List Bullet')

doc.add_page_break()

# Full tweet archive (newest first)
doc.add_heading("Complete Tweet Archive", 1)

for idx, tweet in enumerate(structured_tweets, 1):
    # Timestamp & metrics
    created = datetime.fromisoformat(tweet["created_at"].replace("Z", "+00:00"))
    date_str = created.strftime("%Y-%m-%d %H:%M UTC")
    metrics = tweet["public_metrics"]
    
    # Tweet heading with date
    heading = doc.add_paragraph(
        f"{idx}. {date_str}",
        style='Heading 3'
    )
    
    # Tweet text
    text_para = doc.add_paragraph(tweet["text"])
    text_para.paragraph_format.left_indent = Inches(0.25)
    
    # Metrics line
    metrics_text = f"❤️ {metrics.get('like_count', 0)} | 🔄 {metrics.get('retweet_count', 0)} | 💬 {metrics.get('reply_count', 0)}"
    metrics_para = doc.add_paragraph(metrics_text, style='Normal')
    metrics_para.runs[0].font.size = Pt(9)
    metrics_para.runs[0].font.italic = True
    metrics_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Spacing
    doc.add_paragraph("")

docx_path = Path("/Users/rk/clawd/raul_tweets.docx")
doc.save(docx_path)
print(f"✓ DOCX saved: {docx_path}")

print("\n✅ Complete! Files ready:")
print(f"  • JSON: {json_path}")
print(f"  • DOCX: {docx_path}")
